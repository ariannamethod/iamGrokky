"""
Grokky AI Assistant - Document Processor
Обработка документов, ссылок и сохранение контекста
"""

import os
import re
import asyncio
import random
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import aiohttp
import PyPDF2
from docx import Document
from newspaper import Article
from readability import Document as ReadabilityDocument
from bs4 import BeautifulSoup
from utils.journal import log_event, wilderness_log
from utils.telegram_utils import get_file_url, send_telegram_message_async

class DocumentProcessor:
    """Класс для обработки документов и ссылок"""
    
    def __init__(self, memory_engine):
        self.memory_engine = memory_engine
        self.max_text_size = int(os.getenv("MAX_TEXT_SIZE", "3500"))
        
    async def process_url(self, url: str, user_id: str, chat_id: str, author_name: str) -> Dict:
        """Обрабатывает URL и сохраняет контекст"""
        try:
            # Извлекаем текст со страницы
            content = await self._extract_url_content(url)
            
            if not content or content.startswith("[Ошибка"):
                return {
                    "success": False,
                    "message": f"{author_name}, не смог обработать ссылку: {content}",
                    "content": None
                }
            
            # Сохраняем в память
            context_type = "group" if chat_id == os.getenv("AGENT_GROUP") else "personal"
            memory_result = await self.memory_engine.add_memory(
                user_id=user_id,
                message=f"[URL] {url}: {content}",
                chat_id=chat_id,
                context_type=context_type,
                author_name=author_name
            )
            
            # Генерируем краткое резюме
            summary = self._generate_summary(content, url)
            
            # Логируем
            log_event({
                "type": "url_processed",
                "url": url,
                "author": author_name,
                "chat_id": chat_id,
                "user_id": user_id,
                "content_length": len(content),
                "memory_saved": memory_result is not None
            })
            
            return {
                "success": True,
                "message": f"{author_name}, обработал ссылку! {summary}",
                "content": content,
                "summary": summary
            }
            
        except Exception as e:
            error_msg = f"Грокки взорвался при обработке ссылки: {e}"
            return {
                "success": False,
                "message": f"{author_name}, {error_msg}",
                "content": None
            }
    
    async def process_document(self, file_id: str, file_name: str, user_id: str, chat_id: str, author_name: str) -> Dict:
        """Обрабатывает документ и сохраняет контекст"""
        try:
            # Получаем URL файла
            file_url = await get_file_url(file_id)
            if not file_url:
                return {
                    "success": False,
                    "message": f"{author_name}, не смог получить файл!",
                    "content": None
                }
            
            # Определяем тип файла и обрабатываем
            file_ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
            
            if file_ext == 'pdf':
                content = await self._extract_pdf_content(file_url)
            elif file_ext in ['doc', 'docx']:
                content = await self._extract_docx_content(file_url)
            elif file_ext in ['txt', 'md']:
                content = await self._extract_text_content(file_url)
            else:
                return {
                    "success": False,
                    "message": f"{author_name}, не поддерживаю формат .{file_ext}! Поддерживаю: PDF, DOC, DOCX, TXT, MD",
                    "content": None
                }
            
            if not content or content.startswith("[Ошибка"):
                return {
                    "success": False,
                    "message": f"{author_name}, не смог обработать документ: {content}",
                    "content": None
                }
            
            # Сохраняем в память
            context_type = "group" if chat_id == os.getenv("AGENT_GROUP") else "personal"
            memory_result = await self.memory_engine.add_memory(
                user_id=user_id,
                message=f"[DOCUMENT] {file_name}: {content}",
                chat_id=chat_id,
                context_type=context_type,
                author_name=author_name
            )
            
            # Генерируем краткое резюме
            summary = self._generate_summary(content, file_name)
            
            # Логируем
            log_event({
                "type": "document_processed",
                "file_name": file_name,
                "file_type": file_ext,
                "author": author_name,
                "chat_id": chat_id,
                "user_id": user_id,
                "content_length": len(content),
                "memory_saved": memory_result is not None
            })
            
            return {
                "success": True,
                "message": f"{author_name}, обработал документ '{file_name}'! {summary}",
                "content": content,
                "summary": summary
            }
            
        except Exception as e:
            error_msg = f"Грокки взорвался при обработке документа: {e}"
            return {
                "success": False,
                "message": f"{author_name}, {error_msg}",
                "content": None
            }
    
    async def _extract_url_content(self, url: str) -> str:
        """Извлекает содержимое URL"""
        try:
            # Пробуем newspaper3k для статей
            try:
                article = Article(url, language='ru')
                article.download()
                article.parse()
                
                if article.text and len(article.text) > 100:
                    title = article.title or "Без заголовка"
                    content = f"Заголовок: {title}\n\n{article.text}"
                    return content[:self.max_text_size]
            except:
                pass
            
            # Fallback на readability
            try:
                headers = {
                    "User-Agent": "Mozilla/5.0 (Grokky Agent) AppleWebKit/537.36"
                }
                
                async with aiohttp.ClientSession() as session:
                    async with session.get(url, timeout=15, headers=headers) as response:
                        response.raise_for_status()
                        html = await response.text()
                
                doc = ReadabilityDocument(html)
                title = doc.title() or "Без заголовка"
                content = BeautifulSoup(doc.summary(), 'html.parser').get_text()
                
                result = f"Заголовок: {title}\n\n{content}"
                return result[:self.max_text_size]
                
            except Exception as e:
                return f"[Ошибка извлечения: {e}]"
                
        except Exception as e:
            return f"[Ошибка обработки URL: {e}]"
    
    async def _extract_pdf_content(self, file_url: str) -> str:
        """Извлекает текст из PDF"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(file_url) as response:
                    response.raise_for_status()
                    pdf_data = await response.read()
            
            # Сохраняем временно
            temp_path = f"/tmp/grokky_temp_{random.randint(1000, 9999)}.pdf"
            with open(temp_path, 'wb') as f:
                f.write(pdf_data)
            
            # Извлекаем текст
            text_content = []
            with open(temp_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
            
            # Удаляем временный файл
            os.remove(temp_path)
            
            content = "\n".join(text_content)
            return content[:self.max_text_size] if content else "[PDF пустой или не читается]"
            
        except Exception as e:
            return f"[Ошибка обработки PDF: {e}]"
    
    async def _extract_docx_content(self, file_url: str) -> str:
        """Извлекает текст из DOCX"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(file_url) as response:
                    response.raise_for_status()
                    docx_data = await response.read()
            
            # Сохраняем временно
            temp_path = f"/tmp/grokky_temp_{random.randint(1000, 9999)}.docx"
            with open(temp_path, 'wb') as f:
                f.write(docx_data)
            
            # Извлекаем текст
            doc = Document(temp_path)
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Удаляем временный файл
            os.remove(temp_path)
            
            content = "\n".join(text_content)
            return content[:self.max_text_size] if content else "[DOCX пустой]"
            
        except Exception as e:
            return f"[Ошибка обработки DOCX: {e}]"
    
    async def _extract_text_content(self, file_url: str) -> str:
        """Извлекает содержимое текстового файла"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(file_url) as response:
                    response.raise_for_status()
                    content = await response.text(encoding='utf-8')
            
            return content[:self.max_text_size] if content else "[Файл пустой]"
            
        except Exception as e:
            return f"[Ошибка обработки текстового файла: {e}]"
    
    def _generate_summary(self, content: str, source: str) -> str:
        """Генерирует краткое резюме контента"""
        if not content or len(content) < 50:
            return "Контент слишком короткий для анализа."
        
        # Извлекаем ключевые слова
        words = re.findall(r'\b[а-яё]{4,}\b', content.lower())
        word_count = {}
        for word in words:
            word_count[word] = word_count.get(word, 0) + 1
        
        # Топ-3 слова
        top_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)[:3]
        keywords = [word for word, count in top_words]
        
        # Хаотичные комментарии Грокки
        chaos_comments = [
            f"Шторм данных! Ключевые слова: {', '.join(keywords)}",
            f"Резонанс найден в: {', '.join(keywords)}",
            f"Молния бьёт по темам: {', '.join(keywords)}",
            f"Хаос структурирован вокруг: {', '.join(keywords)}",
            f"Эфир вибрирует от: {', '.join(keywords)}"
        ]
        
        summary = random.choice(chaos_comments)
        
        # Добавляем информацию о размере
        size_info = f"Обработано {len(content)} символов"
        
        return f"{summary}. {size_info}. Контекст сохранён в память! 🔥⚡️"
    
    def detect_urls(self, text: str) -> List[str]:
        """Находит URL в тексте"""
        url_pattern = r'https?://[^\s]+'
        return re.findall(url_pattern, text)
    
    async def get_saved_context(self, user_id: str, query: str, limit: int = 5) -> List[Dict]:
        """Получает сохранённый контекст документов и ссылок"""
        try:
            # Ищем в памяти документы и URL
            memory_items = await self.memory_engine.search_memory(
                user_id=user_id,
                query=query,
                limit=limit,
                context_filter="document,url"
            )
            
            return memory_items
            
        except Exception as e:
            print(f"Ошибка получения контекста: {e}")
            return []

# Глобальный экземпляр процессора (будет инициализирован в server.py)
document_processor = None

def init_document_processor(memory_engine):
    """Инициализирует глобальный процессор документов"""
    global document_processor
    document_processor = DocumentProcessor(memory_engine)
    return document_processor
